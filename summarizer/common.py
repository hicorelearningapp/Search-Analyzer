import os
from config import Config
from docx import Document

class SummarizerPipeline:
    def __init__(self):
        pass

    def run(self, text: str, doc_type: str, pages: int, label: str, outfile_name: str, download: bool = False):
        try:
            # Generate document (example with python-docx)
            document = Document()
            document.add_heading(doc_type, 0)
            document.add_paragraph(f"Topic: {label}")
            document.add_paragraph(text)

            # Always save file to /reports
            if not os.path.exists(Config.REPORTS_DIR):
                os.makedirs(Config.REPORTS_DIR)

            outfile_path = os.path.join(Config.REPORTS_DIR, f"{outfile_name}.docx")
            document.save(outfile_path)

            # If download requested, return a public link
            if download:
                return {
                    "status": "success",
                    "file_url": f"/reports/{outfile_name}.docx"
                }

            # Otherwise return inline summary only
            return {
                "status": "success",
                "summary": text,
                "doc_type": doc_type,
                "pages": pages,
                "label": label
            }

        except Exception as e:
            return {"status": "error", "message": str(e)}
