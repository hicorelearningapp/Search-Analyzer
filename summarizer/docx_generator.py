from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxFormatter:
    """
    Formatter for creating neatly styled, professional DOCX documents.
    """

    def __init__(self, filename: str = "output.docx"):
        self.filename = filename
        self.document = Document()
        self._set_base_styles()

    def _set_base_styles(self):
        """Configure global and heading styles."""
        # Normal style
        normal = self.document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Heading styles
        self._set_heading_style("Heading 1", size=16, bold=True)
        self._set_heading_style("Heading 2", size=14, bold=True)
        self._set_heading_style("Heading 3", size=12, bold=True)

    def _set_heading_style(self, style_name: str, size: int, bold: bool = True):
        """Helper to define heading style properties."""
        style = self.document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold

    def add_title(self, title: str):
        """Add main title as Heading 1."""
        self.document.add_heading(title, level=1)

    def add_metadata(self, metadata: dict):
        """Add metadata section with Heading 2."""
        if metadata:
            self.document.add_heading("Metadata", level=2)
            for key, value in metadata.items():
                self.document.add_paragraph(f"{key}: {value}", style="Normal")

    def add_sections(self, headings: list, content: str):
        """
        Add sections based on provided headings.
        If headings are given, split content into sections.
        """
        if not headings:
            self.document.add_paragraph(content, style="Normal")
            return

        self.document.add_heading("Content", level=2)

        paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
        for i, heading in enumerate(headings):
            self.document.add_heading(heading, level=3)
            if i < len(paragraphs):
                self.document.add_paragraph(paragraphs[i], style="Normal")

    def save(self):
        """Save the document to the given filename."""
        self.document.save(self.filename)
        print(f"✅ DOCX saved as {self.filename}")


class SummaryDocxBuilder:
    """
    Builds a DOCX document from a summarizer output dictionary.
    """

    def __init__(self, summary: dict, filename: str = "output.docx"):
        self.summary = summary
        self.formatter = DocxFormatter(filename)

    def build(self):
        """Build the DOCX document from the summary structure."""
        self.formatter.add_title(self.summary.get("document_type", "Summary Document"))
        self.formatter.add_metadata(self.summary.get("metadata", {}))
        self.formatter.add_sections(
            self.summary.get("headings", []),
            self.summary.get("content", "")
        )
        self.formatter.save()
