# summarizer/docx_generator.py
from docx import Document


class SummaryDocxBuilder:
    def __init__(self, summary: dict, filepath: str):
        self.summary = summary
        self.filepath = filepath
        self.doc = Document()

    def build(self):
        # Grab the generated content directly
        content = self.summary.get("content", "")

        # Split into lines for parsing
        lines = content.splitlines()

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue  # skip empty lines

            # Count leading hashtags
            if stripped.startswith("#"):
                level = len(stripped.split()[0])  # number of '#' in first token
                heading_text = stripped[level:].strip()

                # Clamp heading level to max Word supports (1-9)
                level = min(level, 9)

                self.doc.add_heading(heading_text, level=level)
            else:
                self.doc.add_paragraph(stripped)

        # Save document
        self.doc.save(self.filepath)
